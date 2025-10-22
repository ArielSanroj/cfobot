"""Report generation utilities."""

from __future__ import annotations

from pathlib import Path
from typing import Dict, Iterable, List, Mapping, Sequence

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document

from .config import AppConfig
from .data_loader import FinancialData
from .processing import BudgetResult, KPIResult


def _build_filename(prefix: str, data: FinancialData, suffix: str) -> Path:
    return Path.home() / "Downloads" / f"{prefix}_{data.current_month.lower()}_2025{suffix}"


def save_consolidated_balance(df: pd.DataFrame, data: FinancialData) -> Path:
    output_path = _build_filename("consolidated_balance", data, ".xlsx")
    df.to_excel(output_path, index=False)
    return output_path


def save_budget_execution(budget: BudgetResult, data: FinancialData) -> Path:
    output_path = _build_filename("presupuesto_ejecutado", data, ".xlsx")
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        budget.summary.to_excel(writer, sheet_name="Ejecutado", index=False)
        if not budget.distribution.empty:
            budget.distribution.reset_index().to_excel(
                writer, sheet_name="Distribución", index=False
            )
    return output_path


def save_kpis(kpis: KPIResult, data: FinancialData) -> Path:
    output_path = _build_filename("kpis_financieros", data, ".xlsx")
    kpis.table.to_excel(output_path, index=False)
    return output_path


def generate_spending_chart(budget: BudgetResult, data: FinancialData) -> Path:
    totals: List[float] = []
    for month in data.months:
        mask = data.eri["Codigo"].str.match(r"^(51|53|61|72|73)[0-9]{4,}", na=False)
        totals.append(abs(float(data.eri.loc[mask, month].sum())))

    plt.figure(figsize=(14, 8))
    bars = plt.bar(data.months, totals, color="skyblue", edgecolor="navy", linewidth=1.2)
    
    # Add value labels on top of bars
    for bar, total in zip(bars, totals):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + max(totals)*0.01,
                f'${total:,.0f}', ha='center', va='bottom', fontsize=9, fontweight='bold')
    
    plt.title(f"Gastos Mensuales - Enero a {data.current_month} 2025", 
              fontsize=16, fontweight='bold', pad=20)
    plt.xlabel("Mes", fontsize=12, fontweight='bold')
    plt.ylabel("Gastos Totales (COP)", fontsize=12, fontweight='bold')
    plt.xticks(rotation=45, fontsize=10)
    plt.yticks(fontsize=10)
    plt.grid(axis="y", linestyle="--", alpha=0.7)
    
    # Format y-axis to show values in millions
    ax = plt.gca()
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x/1e6:.1f}M'))
    
    save_path = _build_filename("monthly_spending", data, ".png")
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    return save_path


def generate_kpi_chart(kpis: KPIResult, data: FinancialData) -> Path:
    plt.figure(figsize=(14, 8))
    bars = plt.bar(kpis.table["KPI"], kpis.table.iloc[:, 1], color="teal", edgecolor="darkgreen", linewidth=1.2)
    
    # Add value labels on top of bars
    for bar, value in zip(bars, kpis.table.iloc[:, 1]):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + max(kpis.table.iloc[:, 1])*0.01,
                f'{value:.2f}', ha='center', va='bottom', fontsize=9, fontweight='bold')
    
    plt.title(f"KPIs Financieros - {data.current_month} 2025", 
              fontsize=16, fontweight='bold', pad=20)
    plt.xlabel("KPI", fontsize=12, fontweight='bold')
    plt.ylabel("Valor", fontsize=12, fontweight='bold')
    plt.xticks(rotation=45, ha="right", fontsize=10)
    plt.yticks(fontsize=10)
    plt.grid(axis="y", linestyle="--", alpha=0.7)
    
    save_path = _build_filename("kpi_dashboard", data, ".png")
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    return save_path


def generate_distribution_pie(budget: BudgetResult, data: FinancialData) -> Path | None:
    if budget.distribution.empty:
        return None

    series = budget.distribution[f"% del Total {data.current_month}"]
    values = series.sort_values(ascending=False)
    top_values = values.head(10)
    remaining = values.iloc[10:].sum()
    if remaining:
        top_values.loc["Otros"] = remaining

    plt.figure(figsize=(14, 10))
    colors = plt.cm.Set3(np.linspace(0, 1, len(top_values)))
    wedges, texts, autotexts = plt.pie(
        top_values.values,
        labels=top_values.index,
        autopct="%1.1f%%",
        colors=colors,
        startangle=90,
        textprops={'fontsize': 10, 'fontweight': 'bold'}
    )
    for autotext in autotexts:
        autotext.set_color("white")
        autotext.set_fontweight("bold")
        autotext.set_fontsize(9)

    plt.title(f"Distribución de Gastos - {data.current_month} 2025", 
              fontsize=18, fontweight="bold", pad=20)
    plt.axis("equal")

    # Enhanced legend with better formatting
    legend_labels = [
        f"{name}: {value:.1f}%"
        for name, value in zip(top_values.index, top_values.values)
    ]
    plt.legend(wedges, legend_labels, title="Gastos", loc="center left", 
               bbox_to_anchor=(1, 0, 0.5, 1), fontsize=10, title_fontsize=12)
    
    # Add total expenses text
    total_expenses = budget.distribution[data.current_month_col].sum()
    plt.figtext(0.5, 0.02, f"Total Gastos: ${total_expenses:,.0f} COP", 
                ha="center", fontsize=12, fontweight="bold")
    
    save_path = _build_filename("distribucion_gastos_pie", data, ".png")
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    return save_path


def generate_category_pie(budget: BudgetResult, data: FinancialData) -> Path | None:
    categories = {
        "Administrativos": budget.gastos_admin,
        "Otros Gastos": budget.gastos_otros,
        "Costos de Venta": budget.costos_venta,
        "Costos de Producción": budget.costos_produccion,
    }
    categories = {k: v for k, v in categories.items() if v > 0}

    if not categories:
        return None

    plt.figure(figsize=(10, 8))
    wedges, texts, autotexts = plt.pie(
        categories.values(),
        labels=categories.keys(),
        autopct="%1.1f%%",
        colors=["#FF9999", "#66B2FF", "#99FF99", "#FFCC99", "#FF99CC"][: len(categories)],
        startangle=90,
    )
    for autotext in autotexts:
        autotext.set_color("white")
        autotext.set_fontweight("bold")

    plt.title(f"Distribución por Categorías de Gastos - {data.current_month} 2025", fontsize=14, fontweight="bold")
    plt.axis("equal")
    total_expenses = sum(categories.values())
    plt.figtext(0.5, 0.02, f"Total Gastos: ${total_expenses:,.0f} COP", ha="center", fontsize=12, fontweight="bold")
    save_path = _build_filename("categorias_gastos_pie", data, ".png")
    plt.tight_layout()
    plt.savefig(save_path, dpi=300)
    plt.close()
    return save_path


def build_ai_enhanced_board_report(
    budget: BudgetResult,
    kpis: KPIResult,
    data: FinancialData,
    ai_insights,
    diferencia: float | None = None,
) -> Path:
    """Build AI-enhanced board report with intelligent insights.
    
    Args:
        budget: Budget execution results
        kpis: KPI results
        data: Financial data
        ai_insights: AI-generated insights
        diferencia: Bank reconciliation difference
        
    Returns:
        Path to generated report
    """
    output_path = _build_filename("informe_junta_ai", data, ".docx")
    doc = Document()
    doc.add_heading(f"Informe Ejecutivo con IA - {data.current_month} 2025", 0)

    # AI-Enhanced Executive Summary
    doc.add_heading("Resumen Ejecutivo con Análisis de IA", level=1)
    doc.add_paragraph(ai_insights.executive_summary)

    # Key AI Insights
    doc.add_heading("Insights Clave de IA", level=1)
    for insight in ai_insights.key_insights:
        doc.add_paragraph(f"• {insight}")

    # Financial Performance with AI Analysis
    ingresos = float(budget.summary.loc[0, f"Actual {data.current_month}"])
    gastos = float(budget.summary.loc[1, f"Actual {data.current_month}"])
    ebitda = float(kpis.metrics.get("EBITDA", 0))
    utilidad_neta = float(kpis.metrics.get("Margen Neto %", 0))
    current_ratio = float(kpis.metrics.get("Current Ratio", 0))

    doc.add_heading("Análisis Financiero Detallado", level=1)
    doc.add_paragraph(
        f"Durante el mes de {data.current_month} 2025, la empresa presentó los siguientes resultados:\n\n"
        f"• Ingresos Totales: ${ingresos:,.0f} COP\n"
        f"• Gastos Totales: ${gastos:,.0f} COP\n"
        f"• EBITDA: ${ebitda:,.0f} COP\n"
        f"• Margen Neto: {utilidad_neta:.2f}%\n"
        f"• Ratio de Liquidez (Current Ratio): {current_ratio:.2f}"
    )

    # AI Budget Analysis
    doc.add_heading("Análisis Presupuestario con IA", level=1)
    doc.add_paragraph(ai_insights.budget_analysis)

    # AI KPI Analysis
    doc.add_heading("Análisis de KPIs con IA", level=1)
    doc.add_paragraph(ai_insights.kpi_analysis)

    # AI Trend Analysis
    doc.add_heading("Análisis de Tendencias con IA", level=1)
    doc.add_paragraph(ai_insights.trend_analysis)

    # AI Risk Assessment
    doc.add_heading("Evaluación de Riesgos con IA", level=1)
    doc.add_paragraph(ai_insights.risk_assessment)

    # AI Recommendations
    doc.add_heading("Recomendaciones Estratégicas con IA", level=1)
    for i, recommendation in enumerate(ai_insights.recommendations, 1):
        doc.add_paragraph(f"{i}. {recommendation}")

    # Traditional KPI table
    doc.add_heading("Indicadores Financieros Clave (KPIs)", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicador'
    hdr_cells[1].text = f'Valor {data.current_month} 2025'
    
    for kpi, valor in kpis.table.values:
        row = table.add_row().cells
        row[0].text = str(kpi)
        row[1].text = str(valor)

    # Add footer with AI generation info
    doc.add_paragraph(f"\n\n---\nInforme generado con IA el {data.current_month} 2025\nSistema CFO Bot v2.0 con Ollama")

    doc.save(output_path)
    return output_path


def build_board_report(
    budget: BudgetResult,
    kpis: KPIResult,
    data: FinancialData,
    diferencia: float | None = None,
) -> Path:
    output_path = _build_filename("informe_junta", data, ".docx")
    doc = Document()
    doc.add_heading(f"Informe para Junta Directiva - {data.current_month} 2025", 0)

    # Enhanced financial summary
    ingresos = float(budget.summary.loc[0, f"Actual {data.current_month}"])
    gastos = float(budget.summary.loc[1, f"Actual {data.current_month}"])
    ebitda = float(kpis.metrics.get("EBITDA", 0))
    utilidad_neta = float(kpis.metrics.get("Margen Neto %", 0))
    current_ratio = float(kpis.metrics.get("Current Ratio", 0))

    doc.add_heading("Resumen Financiero Ejecutivo", level=1)
    doc.add_paragraph(
        f"Durante el mes de {data.current_month} 2025, la empresa presentó los siguientes resultados:\n\n"
        f"• Ingresos Totales: ${ingresos:,.0f} COP\n"
        f"• Gastos Totales: ${gastos:,.0f} COP\n"
        f"• EBITDA: ${ebitda:,.0f} COP\n"
        f"• Margen Neto: {utilidad_neta:.2f}%\n"
        f"• Ratio de Liquidez (Current Ratio): {current_ratio:.2f}"
    )

    # Budget execution analysis
    doc.add_heading("Análisis de Ejecución Presupuestaria", level=1)
    ejecutado_ingresos = float(budget.summary.loc[0, "% Ejecutado"])
    ejecutado_gastos = float(budget.summary.loc[1, "% Ejecutado"])
    
    doc.add_paragraph(
        f"La ejecución presupuestaria del mes muestra:\n\n"
        f"• Ingresos: {ejecutado_ingresos:.1f}% del presupuesto mensual\n"
        f"• Gastos: {ejecutado_gastos:.1f}% del presupuesto mensual"
    )

    # Detailed KPI table
    doc.add_heading("Indicadores Financieros Clave (KPIs)", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicador'
    hdr_cells[1].text = f'Valor {data.current_month} 2025'
    
    for kpi, valor in kpis.table.values:
        row = table.add_row().cells
        row[0].text = str(kpi)
        row[1].text = str(valor)

    # Expense breakdown
    doc.add_heading("Desglose de Gastos por Categoría", level=1)
    gastos_admin = float(budget.summary.loc[2, f"Actual {data.current_month}"])
    gastos_otros = float(budget.summary.loc[3, f"Actual {data.current_month}"])
    costos_venta = float(budget.summary.loc[4, f"Actual {data.current_month}"])
    costos_prod = float(budget.summary.loc[5, f"Actual {data.current_month}"])

    doc.add_paragraph(
        f"• Gastos Administrativos: ${gastos_admin:,.0f} COP\n"
        f"• Gastos Otros: ${gastos_otros:,.0f} COP\n"
        f"• Costos de Venta: ${costos_venta:,.0f} COP\n"
        f"• Costos de Producción: ${costos_prod:,.0f} COP"
    )

    # Bank reconciliation
    doc.add_heading("Conciliación Bancaria", level=1)
    if diferencia is not None:
        doc.add_paragraph(
            f"Se identificó una diferencia de ${diferencia:,.0f} COP en la conciliación bancaria, "
            "posiblemente relacionada con consignaciones no acreditadas al cierre del mes."
        )
    else:
        doc.add_paragraph("No se pudo determinar la diferencia de conciliación bancaria.")

    # Enhanced recommendations
    doc.add_heading("Recomendaciones Estratégicas", level=1)
    
    recomendaciones = []
    
    # Liquidity recommendations
    if current_ratio < 1.5:
        recomendaciones.append(
            f"• URGENTE: Mejorar la liquidez inmediata. El Current Ratio de {current_ratio:.2f} "
            "está por debajo del mínimo recomendado de 1.5. Considerar estrategias de recaudo "
            "o financiamiento a corto plazo."
        )
    else:
        recomendaciones.append(
            f"• La liquidez se mantiene en niveles adecuados (Current Ratio: {current_ratio:.2f})."
        )
    
    # Profitability recommendations
    if utilidad_neta < 5:
        recomendaciones.append(
            f"• Revisar la estructura de costos. El margen neto de {utilidad_neta:.2f}% "
            "está por debajo de los estándares de la industria. Analizar costos de venta y gastos operativos."
        )
    
    # Budget execution recommendations
    if ejecutado_gastos > 100:
        recomendaciones.append(
            f"• Control de gastos: Se superó el presupuesto en {ejecutado_gastos-100:.1f}%. "
            "Implementar controles más estrictos en la aprobación de gastos."
        )
    
    if ejecutado_ingresos < 80:
        recomendaciones.append(
            f"• Revisar estrategias de ventas. Los ingresos representan solo {ejecutado_ingresos:.1f}% "
            "del presupuesto. Evaluar canales de venta y estrategias comerciales."
        )
    
    # General recommendations
    recomendaciones.extend([
        "• Monitorear mensualmente los KPIs financieros para detectar tendencias tempranas.",
        "• Implementar un sistema de alertas automáticas para desviaciones presupuestarias significativas.",
        "• Revisar trimestralmente la estructura de costos para optimizar la rentabilidad."
    ])
    
    for rec in recomendaciones:
        doc.add_paragraph(rec)

    # Add footer with generation info
    doc.add_paragraph(f"\n\n---\nInforme generado automáticamente el {data.current_month} 2025\nSistema CFO Bot v1.0")

    doc.save(output_path)
    return output_path


def extract_caratula_difference(data: FinancialData) -> float | None:
    try:
        diferencia_row = data.caratula["Column_0"].astype(str).str.contains("Diferencia", case=False, na=False)
        values = data.caratula.loc[diferencia_row, "Column_1"]
        if values.empty:
            return None
        return float(values.iloc[0])
    except Exception:
        return None


def generate_all_figures(
    budget: BudgetResult,
    kpis: KPIResult,
    data: FinancialData,
) -> List[Path]:
    figures: List[Path] = []
    figures.append(generate_spending_chart(budget, data))
    figures.append(generate_kpi_chart(kpis, data))

    distribution_chart = generate_distribution_pie(budget, data)
    if distribution_chart:
        figures.append(distribution_chart)

    category_chart = generate_category_pie(budget, data)
    if category_chart:
        figures.append(category_chart)

    return figures
